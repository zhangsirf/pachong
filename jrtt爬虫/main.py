#!/usr/bin/env python
# -*- coding: UTF-8 -*-
'''
@Project ：jrttpc
@File ：main.py
@IDE  ：PyCharm 
@Author ：Mr. zhang
@Date ：2023/2/21 下午5:51 
'''
import time
import requests
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document

def get_driver(chrome_driver='./chromedriver.exe',):
    driver = webdriver.Chrome(chrome_driver)
    # driver = webdriver.Chrome(ChromeDriverManager().install())
    # driver.maximize_window()
    # Access requests via the `requests` attribute
    return driver


def start_driver():
    # options = webdriver.ChromeOptions()
    # options.add_argument('start-maximized')
    # options.add_argument('--headless')
    # options.add_experimental_option('excludeSwitches', ['enable-automation'])
    # options.add_experimental_option('useAutomationExtension', False)
    # _driver = webdriver.Chrome(options=options)
    _driver = get_driver()
    with open('./stealth.min.js') as f:
        js = f.read()
    _driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {'source': js})
    return _driver


def check():
    while True:
        soup = BeautifulSoup(driver.page_source, 'lxml')
        article = soup.find_all('div', class_='article-content')
        if article: return
        time.sleep(0.5)


def save_article(link):
    driver.get(link)
    check()
    soup = BeautifulSoup(driver.page_source, 'lxml')
    title = soup.find('h1').string
    article = soup.find('article', class_='syl-article-base')
    images = article.find_all('img')
    like = soup.find('div', class_='detail-like').get('aria-label').replace('点赞', '')
    if images:
        document.add_paragraph(title)
        for image in images:
            aa = image.get('src')
            split_ = str(article).split(str(image))
            string = split_[0]
            article = split_[-1]
            s = BeautifulSoup(string, 'lxml')
            p_list = s.find_all('p')
            for p in p_list:
                tx = p.get_text().strip()
                document.add_paragraph(tx)
            # content = requests.get(image.get('src')).content
            # file_path = f'./img/{time.time()}.jpg'
            # print(file_path)
            # open(file_path, 'wb').write(content)
            res = requests.get(aa)
            file_path = f'./img/{time.time()}.jpg'
            with open(file_path, "wb") as f:
                f.write(res.content)
            try:
                document.add_picture(file_path)
            except:
                document.add_paragraph(aa)
        s = BeautifulSoup(article, 'lxml')
        p_list = s.find_all('p')
        for p in p_list:
            tx = p.get_text().strip()
            document.add_paragraph(tx)
    else:
        p_list = article.find_all('p')
        for p in p_list:
            tx = p.get_text().strip()
            document.add_paragraph(tx)
    document.save('./article/' + title + '--点赞量:' + like + '.docx')


if __name__ == '__main__':
    driver = start_driver()
    document = Document()
    link = ('https://www.toutiao.com/article/7123003396730274307/?channel=&source=search_tab')
    save_article(link=link)
